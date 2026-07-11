from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.infrastructure.documents.md_blocks import parse_blocks


def markdown_to_docx(markdown_text: str, path: Path) -> None:
    document = Document()
    for block in parse_blocks(markdown_text):
        if block.kind == "heading1":
            document.add_heading(block.text, level=1)
        elif block.kind == "heading2":
            document.add_heading(block.text, level=2)
        elif block.kind == "heading3":
            document.add_heading(block.text, level=3)
        elif block.kind == "bullet":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "numbered":
            document.add_paragraph(block.text, style="List Number")
        elif block.kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            document.add_paragraph(block.text)
    document.save(str(path))
